from docx import Document
from pathlib import Path

path = Path(r"c:/Users/Alexandre/projet/formulaireTestAccueil/S-OG-FO-003-ORGA_Test Accueil SSE_Goron_systemes.docx")

# Fallback if filename includes non-ASCII; list and pick the first .docx
if not path.exists():
    candidates = list(Path(r"c:/Users/Alexandre/projet/formulaireTestAccueil").glob("*.docx"))
    if not candidates:
        raise FileNotFoundError("No .docx found in workspace root")
    path = candidates[0]

print(f"Using: {path}")

doc = Document(path)

print("PARAGRAPHS:")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"{i:03d}: {text}")

print("\nTABLES:")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        print(f"  Row {ri}: {cells}")
